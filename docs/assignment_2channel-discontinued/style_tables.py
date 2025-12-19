#!/usr/bin/env python3
"""
Post-process DOCX tables to apply professional styling.
Adds alternating row shading, borders, and Helvetica font.

Usage:
    python style_tables.py <docx_file>
"""

import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Styling constants
FONT_NAME = "Helvetica Neue"
FONT_SIZE = 11
HEADER_BG_COLOR = "E8E8E8"  # Light gray header
HEADER_TEXT_COLOR = RGBColor(0, 0, 0)  # Black text
BORDER_COLOR = "CCCCCC"  # Light gray borders


def set_cell_shading(cell, color_hex):
    """Set background color of a cell."""
    shading_elm = parse_xml(
        f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'
    )
    cell._tc.get_or_add_tcPr().append(shading_elm)


def set_cell_borders(cell, color="CCCCCC", size="4"):
    """Set borders on a cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcBorders = parse_xml(
        f'''<w:tcBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="{size}" w:color="{color}"/>
            <w:left w:val="single" w:sz="{size}" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="{size}" w:color="{color}"/>
            <w:right w:val="single" w:sz="{size}" w:color="{color}"/>
        </w:tcBorders>'''
    )
    tcPr.append(tcBorders)


def style_cell_text(cell, font_name=FONT_NAME, font_size=FONT_SIZE, bold=False, color=None):
    """Style text within a cell."""
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            run.font.bold = bold
            if color:
                run.font.color.rgb = color
        paragraph.style.font.name = font_name


def add_cell_padding(cell):
    """Add padding to cell."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    tcMar = parse_xml(
        f'''<w:tcMar {nsdecls("w")}>
            <w:top w:w="60" w:type="dxa"/>
            <w:left w:w="100" w:type="dxa"/>
            <w:bottom w:w="60" w:type="dxa"/>
            <w:right w:w="100" w:type="dxa"/>
        </w:tcMar>'''
    )
    tcPr.append(tcMar)


def style_table(table):
    """Apply professional styling to a single table."""
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_borders(cell, BORDER_COLOR)
            add_cell_padding(cell)

            if row_idx == 0:
                # Header row - gray background, bold text
                set_cell_shading(cell, HEADER_BG_COLOR)
                style_cell_text(cell, bold=True, color=HEADER_TEXT_COLOR)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
            else:
                style_cell_text(cell)


def style_docx_tables(docx_path: Path) -> bool:
    """Style all tables in a DOCX file."""
    try:
        doc = Document(str(docx_path))

        table_count = len(doc.tables)
        if table_count == 0:
            print("ℹ️  No tables found in document")
            return True

        print(f"📊 Found {table_count} table(s)")

        for i, table in enumerate(doc.tables, 1):
            style_table(table)
            print(f"   ✓ Styled table {i}/{table_count}")

        doc.save(str(docx_path))
        print(f"✅ Tables styled successfully!")
        return True

    except Exception as e:
        print(f"❌ Error styling tables: {e}")
        return False


def main():
    if len(sys.argv) < 2:
        print("Usage: python style_tables.py <docx_file>")
        sys.exit(1)

    docx_path = Path(sys.argv[1])

    if not docx_path.exists():
        print(f"❌ File not found: {docx_path}")
        sys.exit(1)

    print(f"\n📄 Styling tables in: {docx_path.name}")

    if style_docx_tables(docx_path):
        print("🎉 Done!")
    else:
        sys.exit(1)


if __name__ == "__main__":
    main()
